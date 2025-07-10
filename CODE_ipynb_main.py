import pandas as pd

# Keyword panel----------------------------------------------------------------------------------
imp_laws=pd.read_excel("crime_data.xlsx")

crime_type_en=imp_laws["Crime Type"]
law_en=imp_laws["Relevant Law/Section"]
crime_type_hi=imp_laws["Crime Type Hindi"]
law_hi=imp_laws["Relevant Law/Section (Hindi)"]

chhattisgarh_districts = ["बालोद", "बलौदाबाजार", "बलरामपुर", "बस्तर", "बेमेतरा", "बिलासपुर", "दंतेवाड़ा", "धमतरी", "दुर्ग",
                          "गौरेला-पेंड्रा-मरवाही", "गरियाबंद", "जांजगीर-चांपा", "जशपुर", "कवर्धा", "कांकेर", "कोरबा", "कोरिया",
                          "महासमुंद", "मुंगेली", "नारायणपुर", "रायगढ़", "रायपुर", "राजनांदगांव", "सुकमा", "सूरजपुर", "सरगुजा",
                          "बीजापुर", "कोंडागांव", "खैरागढ़-छुईखदान-गंडई", "मोहला-मानपुर-अंबागढ़ चौकी", "सारंगढ़-बिलाईगढ़", "मनीन्द्रगढ़-चिरमिरी-भरतपुर"]


crime_severity_en = {
    'Murder': 10,
    'Attempt to Murder': 8,
    'Rape': 9,
    'Robbery': 7,
    'Assault': 5,
    'Theft': 3,
    'Other': 1
}
crime_severity_hi = {
    'हत्या': 10,
    'हत्या का प्रयास': 8,
    'बलात्कार': 9,
    'डकैती': 7,
    'हमला': 5,
    'चोरी': 3,
    'गैंग': 8,             
    'गैंगस्टर': 8,
    'सामूहिक अपराध': 7,    
    'बंदी': 6,              
    'फिरौती': 7,            
    'हत्या की कोशिश': 8,
    'अन्य': 1
}

crime_severity_sorted = {
    'शराब तस्करी': 6,  # Excise
    'जुआ': 2,  # Gambling
    'गौ हत्या': 6,  # Cow related
    'गौ तस्करी': 5,  # Cow related
    'नशीली दवाओं की तस्करी': 9,  # Drug related illegal items
    'हथियारों की तस्करी': 9,  # Drug related illegal items
    'मादक पदार्थ तस्करी': 9,  # Drug related illegal items
    'अवैध हथियार रखना': 8,  # Drug related illegal items
    'एनडीपीएस': 8,  # Drug related illegal items
    'मादक पदार्थ': 7,  # Drug related illegal items
    'वन्यजीव तस्करी': 7,  # Drug related illegal items
    'अवैध उत्पादों की तस्करी': 7,  # Drug related illegal items
    'फेंटेनिल': 4,  # Drug related illegal items
    'एमडी ड्रग': 4,  # Drug related illegal items
    'मेफेड्रोन': 4,  # Drug related illegal items
    'कोकीन': 4,  # Drug related illegal items
    'ब्राउन शुगर': 4,  # Drug related illegal items
    'हेरोइन': 4,  # Drug related illegal items
    'गांजा': 4,  # Drug related illegal items
    'धार्मिक उन्माद': 7,  # Offence related to Religion
    'नकली नोट': 7,  # Offence related to currency note
    'आतंकवाद': 10,  # Offence against
    'देशद्रोह': 10,  # Offence against
    'माओवाद': 9,  # Offence against
    'नक्सली': 9,  # Offence against
    'नक्सलवाद': 9,  # Offence against
    'पुनर्वास नीति': 2,  # Offence against
    'साइबर अपराध': 6,  # Digital offences
    'डिजिटल गिरफ्तारी': 2,  # Digital offences
    'साइबर धोखाधड़ी': 6,  # Digital offences
    'डकैती': 9,  # Property offences
    'लूट': 8,  # Property offences
    'गृहभेदन': 6,  # Property offences
    'धोखाधड़ी': 6,  # Property offences
    'चोरी': 5,  # Property offences
    'ठगी': 5,  # Property offences
    'सांप्रदायिक हत्या': 6,  # Property offences
    'हत्या': 10,  # Body related crime
    'बलात्कार': 10,  # Body related crime
    'दुष्कर्म': 10,  # Body related crime
    'मानव तस्करी': 10,  # Body related crime
    'नाबालिग से दुष्कर्म': 10,  # Body related crime
    'हत्या का प्रयास': 9,  # Body related crime
    'बलात्कार का प्रयास': 9,  # Body related crime
    'अपहरण': 8,  # Body related crime
    'घरेलू हिंसा': 4,  # Body related crime
    'आत्महत्या के लिए उकसाना': 4,  # Body related crime
    'छेड़छाड़': 6,  # Body related crime
    'मारपीट': 6,  # Body related crime
    'पत्नि पर अत्याचार': 7,  # Body related crime
    'धमकी देना': 5,  # Body related crime
    'ईनामी': 0,  # 
    'अवैध परिवहन': 0,  # 
    'बिक्री': 0,  # 
    'आत्मसमर्पण': 0  # 
}

weapons_keywords = ['पिस्तौल', 'रिवॉल्वर', 'बंदूक', 'चाकू', 'तलवार', 'हथियार']

score_match = [
    "व्यापार",
    "एक से अधिक व्यक्तियों की संलिप्तता",
    "वाहनों का उपयोग",
    "हथियारों का उपयोग",
    "पुलिस अधिकारियों की संलिप्तता",
    "प्रमुख व्यक्तियों की संलिप्तता",
    "राजनीतिक व्यक्तियों की संलिप्तता",
    "अंतरराष्ट्रीय संबंध",
    "आपराधिक गिरोह की संलिप्तता",
    "अंतर्राज्यीय गिरोह की संलिप्तता",
    "राजनीतिक दलों से संबंधित",
    "प्रमुख व्यक्तियों से संबंधित",
    "पाँच से अधिक बच्चों से संबंधित",
    "पाँच से अधिक व्यक्तियों से संबंधित",
    "अवैध/नकली शराब से संबंधित",
    "सरकारी कार्यालयों और संपत्तियों की संलिप्तता",
    "विशेष अवसर से संबंधित",
    "महिलाओं से संबंधित",
    "अनुसूचित जाति और अनुसूचित जनजाति से संबंधित",
    "आरोपी के विरुद्ध पहले से कोई सजा, FIR या चार्जशीट",
    "आरोपी या दोषी पर इनाम घोषित था",
    "पीड़ित पहले भी अपराध का शिकार रह चुका था",
    "साइबर डोमेन से संबंधित"
]

occasion= {
    "दीवाली": 10,
    "होली": 10,
    "ईद": 9,
    "क्रिसमस": 8,
    "रक्षा बंधन": 8,
    "दशहरा": 9,
    "ईद-उल-अजहा (बकरीद)": 8,
    "नवरात्रि": 8,
    "गणेश चतुर्थी": 9,
    "मकर संक्रांति": 7,
    "लोहड़ी": 6,
    "बैसाखी": 6,
    "महाशिवरात्रि": 7,
    "गुरु नानक जयंती": 6,
    "रामनवमी": 6,
    "ईद मिलाद-उन-नबी": 5,
    "जन्माष्टमी": 7,
    "बसंत पंचमी": 6,
    "गुड फ्राइडे": 5,
    "बुद्ध पूर्णिमा": 5,
    "महावीर जयंती": 5,
    "शब-ए-बरात": 4,
    "ओणम": 6,
    "पोंगल": 6,
    "तुलसी विवाह": 3,
    "छठ पूजा": 7,
    "इंटरनेशनल वीमन्स डे": 4,
    "स्वतंत्रता दिवस": 9,
    "गणतंत्र दिवस": 9,
    "गांधी जयंती": 7,
    "शहीद दिवस": 5,
    "शादी": 10,
    "सगाई": 6,
    "गृह प्रवेश": 6,
    "उपनयन संस्कार": 5,
    "नामकरण": 4,
    "अंतिम संस्कार/तेरहवीं": 8,
    "धार्मिक यात्रा/तीर्थ यात्रा": 6,
    "राजनीतिक रैली": 7,
    "स्कूल का कार्यक्रम": 5,
    "कॉलेज फेस्ट": 6,
    "लोकल मेला": 5,
    "खेलकूद का आयोजन": 6
}

# news extraction & NER -------------------------------------------------------------

import re
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline

# Load NER model
model_name = "ai4bharat/IndicNER"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForTokenClassification.from_pretrained(model_name)
ner_pipeline = pipeline("ner", model=model, tokenizer=tokenizer, aggregation_strategy="simple")

def extract_date(text):
    # Convert Hindi digits to Western numerals
    hindi_to_num = str.maketrans("०१२३४५६७८९", "0123456789")
    text = text.translate(hindi_to_num)

    # Match common date formats: dd/mm/yyyy, dd-mm-yyyy, dd.mm.yyyy, dd/mm/yy etc.
    match = re.search(r'\b\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4}\b', text)
    return match.group() if match else "N/A"

def extract_district(text):
    possible_districts = []

    # Match patterns like 'जिला <district>'
    match = re.search(r'जिला\s+([^\n।:,]*)', text)
    if match:
        possible_districts.append(match.group(1).strip())

    # Match 'चौकी <...> जिला <district>'
    match_alt = re.search(r'चौकी\s+[^\n।:,]*\s+जिला\s+([^\n।:,]*)', text)
    if match_alt:
        possible_districts.append(match_alt.group(1).strip())

    # Match implicit mentions like 'रायपुर पुलिस', 'रायपुर रेंज', etc.
    for known_district in chhattisgarh_districts:
        if re.search(fr'\b{known_district}\b.*पुलिस', text) or re.search(fr'पुलिस.*\b{known_district}\b', text):
            return known_district

    # Validate against known districts from earlier matches
    for district in possible_districts:
        for known_district in chhattisgarh_districts:
            if known_district in district:
                return known_district

    return "N/A"

def extract_police_station(text):
    # Try both 'थाना' and 'चौकी'
    match = re.search(r'(थाना|चौकी)\s+([^\n।:,]*)', text)
    return f'{match.group(1)} {match.group(2).strip()}' if match else "N/A"

def extract_fir_number(text):
    match = re.search(r'(अपराध क्रमांक|FIR|एफआईआर)\s*[:\-]?\s*([0-9\/\-]+)', text)
    return match.group(2) if match else "N/A"

def extract_complainant(text):
    match = re.search(r'प्रार्थी\s+([^\s,।\n]+(?:\s+[^\s,।\n]+)?)', text)
    return match.group(1) if match else "N/A"

def extract_accused(text):
    match = re.search(r'गिरफ्तार आरोपी\s*-\s*([^\n]*)', text)
    if match:
        name = match.group(1).split('पिता')[0].strip()
        return name
    else:
        # fallback using NER
        ner_results = ner_pipeline(text)
        persons = [ent['word'] for ent in ner_results if ent['entity_group'] == 'PER']
        return persons[0] if persons else "N/A"

def extract_weapons(text):
    weapons_keyword = weapons_keywords
    found = [word for word in weapons_keyword if word in text]
    return ', '.join(found) if found else "N/A"

def extract_crime_type(text):
    for crime in crime_severity_sorted.keys():
        if crime in text:
            return crime
    return "N/A" 

def extract_score(text):
    score=0
    for i in score_match:
        if i in text:
            score += 1
    
    for i in occasion.keys():
        if i in text:
            score += occasion[i]
               
    word_to_number = {
            "शून्य": 0, "एक": 1, "दो": 2, "तीन": 3, "चार": 4, "पांच": 5,
            "छह": 6, "सात": 7, "आठ": 8, "नौ": 9, "दस": 10, "ग्यारह": 11,
            "बारह": 12, "तेरह": 13, "चौदह": 14, "पंद्रह": 15, "सोलह": 16,
            "सत्रह": 17, "अठारह": 18, "उन्नीस": 19, "बीस": 20
        }

    def convert_hindi_numerals(s):
        return s.translate(str.maketrans('०१२३४५६७८९', '0123456789'))

    def extract_number_from_words(w):
        return word_to_number.get(w.strip(), 0)

    crime = extract_crime_type(text)
    if crime in crime_severity_sorted:
        score += crime_severity_sorted[crime]

        if crime == "साइबर धोखाधड़ी":
            # लाख
            lakh = re.search(r'([\d०१२३४५६७८९]+|[अ-ह]+)\s*लाख', text)
            if lakh:
                lakh_word = lakh.group(1)
                if lakh_word.isdigit() or all(ch in '०१२३४५६७८९' for ch in lakh_word):
                    lakh_value = convert_hindi_numerals(lakh_word)
                else:
                    lakh_value = extract_number_from_words(lakh_word)
                score += int(lakh_value) * 2
            # करोड़
            crore = re.search(r'([\d०१२३४५६७८९]+|[अ-ह]+)\s*करोड़', text)
            if crore:
                crore_word = crore.group(1)
                if crore_word.isdigit() or all(ch in '०१२३४५६७८९' for ch in crore_word):
                    crore_value = convert_hindi_numerals(crore_word)
                else:
                    crore_value = extract_number_from_words(crore_word)
                score += int(crore_value) * 20

    return score
    
def extract_summary_string(text):
    info = {
        "Date": extract_date(text),
        "District": extract_district(text),
        "Police Station": extract_police_station(text),
        "FIR_no": extract_fir_number(text),
        "Complainant": extract_complainant(text),
        "Accused": extract_accused(text),
        "Weapons": extract_weapons(text),
        "Type of Crime": extract_crime_type(text)
    }
    # summary = "\n".join([f"{k}: {v}" for k, v in info.items()])
    
    heading = f"🗞️ {extract_district(text)},  {extract_police_station(text)}___{extract_date(text)}___प्राथमिकी संख्या {extract_fir_number(text)}"
    summary = (
        f"{extract_accused(text)} को {extract_crime_type(text)} के मामले में आरोपी बनाया गया है। "
        f"इस प्रकरण में शिकायतकर्ता का नाम {extract_complainant(text)} है। "
        # f"यह रिपोर्ट CCTNS रिपोर्ट में भी उल्लेखित है।"
    ) 
    
    return heading,summary

# Extract News from image --------------------------------------------------------------

from PIL import Image
import pytesseract
import cv2
import re
import os

# === TESSERACT CONFIGURATION ===
pytesseract.pytesseract.tesseract_cmd = '/opt/homebrew/bin/tesseract'
custom_config = r'--oem 3 --psm 6 -l hin'

# === TEXT CLEANING FUNCTION ===
def clean_text(text):
    text = re.sub(r'[^ऀ-ॿa-zA-Z0-9\s।,!?%():\-–—"“”‘’\'\n]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# === EXTRACT FULL TEXT FROM IMAGE ===
def extract_text_from_image(image_path):
    image = cv2.imread(image_path)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    text = pytesseract.image_to_string(gray, config=custom_config)
    return clean_text(text)

# === MAIN FUNCTION ===
def extract_news_bodies_from_images(image_folder="news_images"):
    image_files = sorted([f for f in os.listdir(image_folder)
                          if f.lower().endswith(('.png', '.jpg', '.jpeg'))])

    all_texts = []

    for img_name in image_files:
        img_path = os.path.join(image_folder, img_name)
        body_text = extract_text_from_image(img_path)
        all_texts.append(body_text)

    return all_texts

news_list = extract_news_bodies_from_images("pdf_images")                   #image folder path   

# Extract News from image --------------------------------------------------------------

all_info_2d_list = []
for news in news_list:
    heading, summary = extract_summary_string(news)

    info_row = [
        extract_police_station(news),
        extract_district(news),
        extract_date(news),
        extract_fir_number(news),
        extract_complainant(news),
        extract_accused(news),
        extract_weapons(news),
        extract_crime_type(news),
        extract_score(news),
        heading,
        summary
    ]
    all_info_2d_list.append(info_row)
    
# CCTNS excel ------------------------------------------------------------- 

df_cctns = pd.read_excel("CCTNS_report.xlsx")
df_cctns.columns = df_cctns.iloc[2]
df_cctns = df_cctns.iloc[3:20].reset_index(drop=True)

def generate_CCTNS_report(record):
    fir_id = record['प्राथमिकी संख्या']
    accused_name = record.get('गिरफ्तार व्यक्ति का नाम', 'नाम उपलब्ध नहीं')
    fir_date = pd.to_datetime(record.get('वास्तविक दिनांक और समय (एफआईआर)'), errors='coerce')
    fir_date_str = fir_date.strftime("%d %B %Y") if pd.notnull(fir_date) else "अज्ञात तिथि"
    
    section = record.get('अधिनियम- धारा', 'अनिर्दिष्ट अधिनियम')
    victim_name = record.get('पीड़ित का नाम', 'नाम उपलब्ध नहीं')
    # victim_type = record.get('पीड़ित का प्रकार (एफ आई आर)', 'अज्ञात')
    # victim_gender = record.get('पीड़ित का लिंग (एफ आई आर)', '—')
    # victim_category = record.get('पीड़ित की श्रेणी', '—')
    
    district= record.get('ज़िला','ज़िला उपलब्ध नहीं')
    police_station=record.get('थाना','थाना उपलब्ध नहीं')

    arrest_date = pd.to_datetime(record.get('गिरफ्तारी की तारीख और समय'), errors='coerce')
    arrest_str = (
        f"{arrest_date.strftime('%d %B %Y')} को गिरफ्तार किया गया।"
        if pd.notnull(arrest_date)
        else "अभी तक गिरफ्तारी नहीं हुई है।"
    )

    
    heading_CCTNS = f"{district},  {police_station}___{fir_date}___प्राथमिकी संख्या {fir_id}"
    summary_CCTNS = (
        f"{accused_name} को {section} के अंतर्गत एक मामला दर्ज किया गया। "
        f"पीड़ित का नाम {victim_name} है "
        f"यह रिपोर्ट अखबारों में भी छपी है।"
    ) 

    return heading_CCTNS, summary_CCTNS

# Gruesome Engine -------------------------------------------------------------

def Gruesome_check(row):
    section = row['अधिनियम- धारा']
    accused_count=len(row['गिरफ्तार व्यक्ति का नाम'].split(','))
    victim_count=len(row['पीड़ित का नाम'].split(','))

    if any(str(law) in section for law in law_hi):
        return row  
    if accused_count>=4:
        return row
    if victim_count>=4:
        return row
    return None  

def news_check(row):
    थाना_name = str(row['थाना']).strip()
    for i in range(len(all_info_2d_list)):
        if थाना_name == all_info_2d_list[i][0]:  # Compare with the first element of each row
            return row , i
    return None, None

# Doc Section -------------------------------------------------------------

complete = []
for idx, row in df_cctns.iterrows():
    score = 0
    
    #CCTNS + News
    selected_row_news, index = news_check(row)
    if selected_row_news is not None and index is not None:
        heading_CCTNS, summary_CCTNS = generate_CCTNS_report(selected_row_news)
        newspaper= (
            f"यह मामला अखबार के अनुसार दिनांक {all_info_2d_list[index][2]} का है "
            f"और शिकायतकर्ता का नाम {all_info_2d_list[index][4]} है "
            f"तथा आरोपी का नाम {all_info_2d_list[index][5]} है। "
            f"यह {all_info_2d_list[index][5]} का मामला है।"
        )
        
        accused_count=len(selected_row_news['गिरफ्तार व्यक्ति का नाम'].split(','))
        victim_count=len(selected_row_news['पीड़ित का नाम'].split(','))
        
        score = all_info_2d_list[index][8] + accused_count + victim_count + 5     # extra 5 for feturing on news 
        
        data=[heading_CCTNS, summary_CCTNS +'\n'+ newspaper, score]
        complete.append(data)
        all_info_2d_list.pop(index)  #pop that news from news list 
        
    #CCTNS
    else:
        selected_row = Gruesome_check(row)
        
        if selected_row is not None:
            heading_CCTNS, summary_CCTNS = generate_CCTNS_report(selected_row)
            
            accused_count=len(selected_row['गिरफ्तार व्यक्ति का नाम'].split(','))
            victim_count=len(selected_row['पीड़ित का नाम'].split(','))
            
            c_s=0             
            for i in range(len(law_hi)):
                if selected_row['अधिनियम- धारा'] == law_hi[i]:
                    if crime_type_hi[i] in  crime_severity_sorted:
                        c_s += crime_severity_sorted[crime_type_hi[i]]
 
            score = c_s + accused_count + victim_count 
            
            data=[heading_CCTNS, summary_CCTNS, score]
            complete.append(data)

#rest of the news
for i in range(len(all_info_2d_list)):
    complete.append([all_info_2d_list[i][9], all_info_2d_list[i][10], all_info_2d_list[i][8]])
    
    
# Creating DOC file -------------------------------------------------------------

from docx import Document
from docx.shared import Pt
from datetime import datetime

doc = Document()

#CCTNS Report    
doc.add_heading("अपराध रिपोर्ट", 0)

# Sort the list by score (descending)
sorted_complete = sorted(complete, key=lambda x: float(x[2]), reverse=True)

for i in range(len(sorted_complete)):
    heading, summary, score = sorted_complete[i]
    
    if score >= 5 :
        doc.add_heading(f"{heading}", level=1)
        doc.add_paragraph(f"{summary}")
        doc.add_paragraph(f"स्कोर: {score}")
        doc.add_paragraph('-' * 40)
        doc.add_paragraph("\n")  # Add a new line for better readability

# Save the document with date and time in the filename
today_str = datetime.today().strftime("%Y-%m-%d_%H-%M-%S") 
filename = f"Complete_result\\Crime_Report_{today_str}.docx"
doc.save(filename)

print(f"Report saved as '{filename}'")
print(doc)